from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/propuesta-postres-beinetti-whatsapp.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 24, 28)
MUTED = RGBColor(91, 99, 110)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
BORDER = "D9E2EC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="4"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def set_run_font(run, size=None, bold=None, color=INK):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_paragraph(doc, text="", style=None, bold=False, color=INK, size=11, before=0, after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 6)
    r = p.add_run(text)
    set_run_font(r, size=16 if level == 1 else 13, bold=True, color=BLUE if level == 1 else DARK_BLUE)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    set_cell_border(cell)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=11.5, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.10
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_two_col_table(doc, rows, widths=(2.15, 4.15), header_fill=None):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(widths[0])
        cells[1].width = Inches(widths[1])
        for idx, cell in enumerate(cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_cell_margins(cell)
            if header_fill and idx == 0:
                set_cell_shading(cell, header_fill)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(label if idx == 0 else value)
            set_run_font(run, size=10.5, bold=(idx == 0), color=DARK_BLUE if idx == 0 else INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_three_col_table(doc, headers, rows, widths=(2.0, 2.0, 2.3)):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.width = Inches(widths[i])
        set_cell_shading(cell, LIGHT_FILL)
        set_cell_border(cell)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            cell.width = Inches(widths[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            r = p.add_run(value)
            set_run_font(r, size=10.2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color in (("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 12, DARK_BLUE)):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def set_header_footer(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = ""
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("Postres Beinetti | Automatización WhatsApp")
    set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = ""
    run = footer.add_run("Propuesta comercial")
    set_run_font(run, size=9, color=MUTED)


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    set_header_footer(doc)

    # Cover / masthead
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("PROPUESTA COMERCIAL")
    set_run_font(r, size=12, bold=True, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Automatización de WhatsApp para Postres Beinetti")
    set_run_font(r, size=24, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("MVP profesional para reducir trabajo manual, ordenar pedidos y preparar la integración futura con Ágora POS.")
    set_run_font(r, size=12.5, color=MUTED)

    add_two_col_table(
        doc,
        [
            ("Cliente", "Postres Beinetti"),
            ("Proyecto", "Automatización de WhatsApp Business y gestión inicial de pedidos"),
            ("Fase propuesta", "Fase 1 - MVP WhatsApp automatizado"),
            ("Fecha", "7 de julio de 2026"),
            ("Inversión inicial", "3.500 €"),
            ("Mantenimiento", "250 €/mes + costes externos"),
        ],
        header_fill=LIGHT_FILL,
    )

    add_callout(
        doc,
        "Recomendación",
        "Empezar con Twilio para acelerar la conexión con WhatsApp, reducir complejidad técnica inicial y lanzar una primera versión útil sin esperar a la integración completa con Ágora POS.",
    )

    add_heading(doc, "1. Situación actual", 1)
    add_paragraph(
        doc,
        "Postres Beinetti recibe una parte importante de sus pedidos y consultas por WhatsApp Business. Este canal funciona, pero actualmente exige demasiada gestión manual: conversaciones repetitivas, pedidos mezclados con consultas, restaurantes que escriben o llaman, tartas especiales que necesitan varios datos y posibles errores humanos al pasar la información a producción.",
    )
    add_paragraph(
        doc,
        "La propuesta no cambia el canal que ya utilizan los clientes. La idea es profesionalizar WhatsApp por detrás, automatizando lo repetitivo y dejando a las personas del equipo solo las conversaciones que realmente necesitan criterio humano.",
    )

    add_heading(doc, "2. Solución propuesta", 1)
    add_paragraph(
        doc,
        "Desarrollar una capa de automatización sobre WhatsApp Business que clasifique conversaciones, responda preguntas frecuentes, dirija pedidos normales a la tienda online, recoja datos de tartas especiales y estructure pedidos de restaurantes.",
    )
    for item in [
        "Pedidos particulares normales: derivación clara a la tienda online/Ágora.",
        "Tartas especiales, comuniones, bodas y cumpleaños: recogida de datos y derivación a revisión humana.",
        "Restaurantes: recogida estructurada del pedido y envío al equipo en formato claro.",
        "Atención al cliente: respuestas automáticas para horarios, tiendas, recogidas, alérgenos y dudas frecuentes.",
        "Derivación a humano cuando el caso sea sensible, complejo o requiera presupuesto final.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3. Por qué Twilio", 1)
    add_paragraph(
        doc,
        "Twilio es el proveedor técnico que conecta WhatsApp Business con nuestro software. Para el cliente final no cambia nada: seguirá escribiendo al WhatsApp de Postres Beinetti. Twilio trabaja por detrás, recibiendo el mensaje y enviándolo al sistema para que el bot o el equipo puedan responder.",
    )
    add_paragraph(doc, "Flujo simplificado:", bold=True, color=DARK_BLUE)
    add_paragraph(
        doc,
        "Cliente escribe por WhatsApp -> Twilio recibe el mensaje -> nuestro sistema procesa la conversación -> el cliente recibe la respuesta en WhatsApp.",
    )
    add_paragraph(
        doc,
        "Elegimos Twilio para esta primera fase porque permite avanzar más rápido, tiene documentación sólida, reduce fricción de configuración y mantiene abierta la posibilidad de escalar o integrar Ágora POS más adelante.",
    )

    add_heading(doc, "4. Stack profesional y escalable", 1)
    add_paragraph(
        doc,
        "El stack se elige para construir una primera versión rápida, pero con base profesional. La prioridad es que el MVP no sea un experimento aislado, sino una base que pueda crecer con nuevas tiendas, más volumen, integración con Ágora POS y futuras capacidades de IA o voz.",
    )
    add_three_col_table(
        doc,
        ("Componente", "Tecnología", "Motivo"),
        [
            ("Mensajería WhatsApp", "Twilio WhatsApp API", "Lanzamiento rápido, proveedor estable y menor complejidad inicial."),
            ("Backend", "Node.js + TypeScript", "Base moderna, mantenible y preparada para integraciones API."),
            ("Base de datos", "PostgreSQL", "Fiable para pedidos, clientes, restaurantes, estados y auditoría."),
            ("Panel interno", "Next.js", "Interfaz web profesional para revisar pedidos y conversaciones."),
            ("IA opcional", "OpenAI API", "Clasificación de mensajes, extracción de datos y resúmenes cuando aporte valor."),
            ("Infraestructura", "Railway/Render/Vercel", "Despliegue ágil, costes razonables y escalabilidad suficiente para MVP."),
        ],
    )

    add_heading(doc, "5. Alcance de la Fase 1", 1)
    for item in [
        "Configuración inicial de Twilio para WhatsApp Business.",
        "Diseño de flujos conversacionales para particulares, tartas especiales, restaurantes y atención al cliente.",
        "Bot inicial con preguntas guiadas y respuestas frecuentes.",
        "Recogida de datos clave para tartas especiales: fecha, tienda, personas, tipo de tarta, temática, referencia visual y observaciones.",
        "Recogida de pedidos de restaurantes y envío al equipo por email o panel básico.",
        "Base de datos para registrar pedidos, estados y conversaciones relevantes.",
        "Panel básico para revisar pedidos entrantes y marcar estados.",
        "Despliegue inicial, pruebas y documentación básica.",
        "Dos semanas de ajustes tras la puesta en marcha.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "6. Fuera de alcance inicial", 1)
    add_paragraph(
        doc,
        "Para mantener el MVP controlado y lanzar rápido, los siguientes puntos se recomiendan como fases posteriores:",
    )
    for item in [
        "Integración completa con Ágora POS.",
        "Agente de voz para llamadas.",
        "Automatización total de presupuestos complejos.",
        "Campañas de marketing por WhatsApp.",
        "CRM avanzado o analítica empresarial completa.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "7. Fases posteriores recomendadas", 1)
    for item in [
        "Fase 2: mejorar panel interno, catálogo editable, estados de pedido y reporting diario de producción.",
        "Fase 3: integrar con Ágora POS cuando esté disponible la documentación/API necesaria.",
        "Fase 4: añadir agente IA más avanzado, multilingüe y con soporte de voz si el volumen lo justifica.",
    ]:
        add_number(doc, item)

    add_heading(doc, "8. Inversión", 1)
    add_two_col_table(
        doc,
        [
            ("Desarrollo Fase 1", "3.500 €"),
            ("Mantenimiento mensual", "250 €/mes"),
            ("Costes externos", "Se facturan aparte según consumo real"),
            ("Ajustes incluidos", "2 semanas posteriores al lanzamiento"),
        ],
        header_fill=LIGHT_FILL,
    )
    add_paragraph(
        doc,
        "Nota fiscal: si el proveedor aplica el régimen de pequeño empresario en Alemania/Kleinunternehmerregelung (§19 UStG), la factura se emitirá sin IVA. Este punto debe confirmarse fiscalmente antes de emitir la factura.",
        size=10.5,
        color=MUTED,
    )

    add_heading(doc, "9. Estimación de costes externos del MVP", 1)
    add_paragraph(
        doc,
        "Los costes externos dependen del volumen real de mensajes, uso de IA y proveedor de hosting. Para una primera versión, se puede trabajar con una estimación prudente de 75-200 €/mes.",
    )
    add_three_col_table(
        doc,
        ("Concepto", "Estimación mensual", "Notas"),
        [
            ("Twilio + WhatsApp", "20-100 €", "Variable según mensajes y tarifas aplicables de WhatsApp/Meta."),
            ("Hosting backend/panel", "20-60 €", "Railway, Render, Vercel o equivalente."),
            ("Base de datos", "0-30 €", "PostgreSQL gestionado; puede empezar en plan bajo."),
            ("IA/OpenAI", "10-50 €", "Solo si se usa para clasificar, resumir o extraer datos."),
            ("Monitorización/logs", "0-20 €", "Herramientas básicas de errores y seguimiento."),
        ],
    )
    add_callout(
        doc,
        "Coste mensual esperado",
        "Para el MVP, una previsión razonable es 250 €/mes de mantenimiento + aproximadamente 75-200 €/mes de costes externos. El coste real se revisará tras las primeras semanas de uso.",
    )

    add_heading(doc, "10. Beneficios esperados", 1)
    for item in [
        "Menos tiempo dedicado a conversaciones repetitivas.",
        "Menos errores al recoger pedidos y datos de recogida.",
        "Separación clara entre pedidos normales, restaurantes, tartas especiales y atención al cliente.",
        "Mejor experiencia para clientes que reciben respuestas rápidas.",
        "Base técnica preparada para integración futura con Ágora POS.",
        "Sistema escalable para nuevas tiendas y mayor volumen.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "11. Impacto operativo esperado", 1)
    add_paragraph(
        doc,
        "El principal retorno del proyecto está en liberar tiempo del equipo y reducir errores operativos. En lugar de que una persona tenga que leer, clasificar y responder manualmente cada conversación desde cero, el sistema filtra lo repetitivo, pide los datos necesarios y deja los casos importantes mejor preparados para revisión humana.",
    )
    add_three_col_table(
        doc,
        ("Área", "Problema actual", "Mejora esperada"),
        [
            (
                "Tiempo del equipo",
                "Muchas conversaciones empiezan con preguntas repetidas sobre horarios, recogida, tiendas, alérgenos o cómo hacer pedidos.",
                "El bot responde automáticamente y deriva solo los casos que necesitan criterio humano.",
            ),
            (
                "Errores en pedidos",
                "Los datos llegan mezclados en mensajes: fecha, hora, tienda, número de personas, sabor o cantidades.",
                "El sistema recoge la información de forma guiada y crea un resumen claro antes de confirmar.",
            ),
            (
                "Restaurantes",
                "Los pedidos por WhatsApp pueden quedar poco estructurados o depender de que alguien los copie correctamente.",
                "Cada pedido se transforma en un formato revisable y listo para enviar al equipo o, más adelante, a Ágora POS.",
            ),
            (
                "Tartas especiales",
                "Cada presupuesto requiere varias preguntas y puede faltar información clave.",
                "El sistema recopila datos mínimos antes de pasar el caso a una persona, reduciendo idas y vueltas.",
            ),
            (
                "Producción",
                "La información puede llegar incompleta o tarde al obrador.",
                "Los pedidos quedan registrados con estado, fecha, tienda y detalles principales.",
            ),
        ],
    )
    add_callout(
        doc,
        "Objetivo de ahorro",
        "Como referencia prudente, el MVP debería reducir una parte importante de las conversaciones manuales repetitivas y evitar errores frecuentes de captura de datos. El ahorro exacto se medirá tras las primeras semanas, comparando volumen de conversaciones, pedidos estructurados y casos derivados a humano.",
    )

    add_heading(doc, "12. Próximos pasos", 1)
    for item in [
        "Confirmar acceso y situación actual de WhatsApp Business/Meta Business.",
        "Definir los flujos exactos de pedidos particulares, tartas especiales y restaurantes.",
        "Recopilar preguntas frecuentes, horarios, tiendas, catálogo base y condiciones de recogida.",
        "Preparar entorno Twilio y primer prototipo funcional.",
        "Probar con conversaciones reales anonimizadas antes del lanzamiento.",
    ]:
        add_number(doc, item)

    add_paragraph(
        doc,
        "La recomendación es empezar con una primera fase controlada, enfocada en ahorrar tiempo y ordenar pedidos. A partir de datos reales de uso, se podrá decidir con más precisión qué automatizar después y cuándo abordar la integración con Ágora POS.",
        bold=True,
        color=DARK_BLUE,
        before=10,
        after=0,
    )

    doc.save(OUT)


if __name__ == "__main__":
    build()
